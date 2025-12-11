import streamlit as st
import os
import fitz
import docx
import pandas as pd
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from openai import OpenAI
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive

# Google Drive Authentication
def load_drive():
    gauth = GoogleAuth()
    gauth.settings.update({
        "client_config": {
            "client_id": st.secrets["google_oauth"]["client_id"],
            "client_secret": st.secrets["google_oauth"]["client_secret"],
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://accounts.google.com/o/oauth2/token",
            "redirect_uris": ["urn:ietf:wg:oauth:2.0:oob", "http://localhost"],
        }
    })
    gauth.LocalWebserverAuth()
    drive = GoogleDrive(gauth)
    return drive
def download_folder(drive, folder_id, local_path):
    if not os.path.exists(local_path):
        os.makedirs(local_path)

    file_list = drive.ListFile({'q': f"'{folder_id}' in parents and trashed=false"}).GetList()

    for file in file_list:
        fname = os.path.join(local_path, file['title'])
        file.GetContentFile(fname)
        print("Downloaded:", fname)

# client = OpenAI(api_key=)   
api_key = st.secrets.get("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# File Extraction :
def extract_pdf(path, filename):
    doc = fitz.open(path)
    pages = []

    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append({
                "id": f"{filename}__page_{i+1}",
                "source": filename,
                "type": "pdf",
                "text": text.strip()
            })
    return pages

def extract_word(path, filename):
    try:
        doc_file = docx.Document(path)
    except:
        return []

    output = []

    # Paragraphs
    paragraphs = [p.text.strip() for p in doc_file.paragraphs if p.text.strip()]
    if paragraphs:
        output.append({
            "id": f"{filename}__paragraphs",
            "source": filename,
            "type": "word",
            "text": "\n".join(paragraphs)
        })

    # Tables
    table_index = 0
    for table in doc_file.tables:
        table_index += 1
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        table_text = "\n".join(rows_text).strip()

        if table_text:
            output.append({
                "id": f"{filename}__table_{table_index}",
                "source": filename,
                "type": "word_table",
                "text": table_text
            })

    return output

def extract_excel(path, filename):
    try:
        sheets = pd.read_excel(path, sheet_name=None)
    except:
        return []

    output = []

    for sheet, df in sheets.items():
        if df.empty:
            continue

        headers = list(df.columns)

        for idx, row in df.iterrows():
            fields = []
            for h, v in zip(headers, row.tolist()):
                v = str(v).strip()
                if v.lower() == "nan" or v == "":
                    continue
                fields.append(f"{h}: {v}")

            if fields:
                row_txt = "\n".join(fields)
                output.append({
                    "id": f"{filename}__sheet_{sheet}__row_{idx}",
                    "source": filename,
                    "type": "excel",
                    "text": row_txt
                })

    return output

def process_dir(folder):
    records = []

    for fname in os.listdir(folder):
        path = os.path.join(folder, fname)
        ext = fname.lower().split(".")[-1]

        if ext == "pdf":
            recs = extract_pdf(path, fname)
        elif ext == "docx":
            recs = extract_word(path, fname)
        elif ext in ["xlsx", "xls"]:
            recs = extract_excel(path, fname)
        else:
            continue

        records.extend(recs)

    return records

# Embeddings & Knoweldge Database
def build_faiss_index(texts):
    model = SentenceTransformer("sentence-transformers/paraphrase-multilingual-mpnet-base-v2") # General Model
    embs = model.encode(texts, show_progress_bar=True)
    embs = np.asarray(embs).astype("float32")

    index = faiss.IndexFlatL2(embs.shape[1])
    index.add(embs)

    return index, model, embs

# RAG Search
def search(query, model, index, ids, texts, k=10):
    q_emb = model.encode([query], convert_to_numpy=True).astype("float32")
    D, I = index.search(q_emb, k)

    results = []
    for rank, idx in enumerate(I[0]):
        results.append({
            "rank": rank + 1,
            "score": float(D[0][rank]),
            "id": ids[idx],
            "text": texts[idx]
        })
    return results

# LLM Prompt
def build_context(results):
    ctx = ""
    for r in results:
        ctx += f"\n\n--- من الملف: {r['id']} ---\n{r['text']}"
    return ctx


def llm_answer(query, results):
    context = build_context(results)

    prompt = f"""
أجب على السؤال التالي اعتمادًا فقط على المحتوى التالي:

السؤال:
{query}

السياق:
{context}

الإجابة:
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "أنت مساعد جامعي دقيق جدًا ولا تذكر معلومات غير موجودة في المستندات."},
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content

# STREAMLIT UI
st.title("SEU Chatbot")

# FAQ_DIR = st.text_input("📁 مسار مجلد FAQ", r"D:\Azzam\Personal_Projects\SEU\filtered_data\FAQ")
# DOCS_DIR = st.text_input("📁 مسار مجلد Docs", r"D:\Azzam\Personal_Projects\SEU\filtered_data\Docs")
FAQ_DIR = st.text_input("📁 مجلد FAQ", "FAQ")
DOCS_DIR = st.text_input("📁 مجلد Docs", "Docs")

if st.button("بناء قاعدة المعرفة"):
    with st.spinner("جاري بناء قاعدة البيانات من الملفات..."):
        with st.spinner("Downloading files from Google Drive..."):
            drive = load_drive()
            download_folder(drive, FAQ_DRIVE_ID, "FAQ")
            download_folder(drive, DOCS_DRIVE_ID, "Docs")
    
        st.success("تم تحميل جميع الملفات من Google Drive!")
        records = process_dir(FAQ_DIR)
        records += process_dir(DOCS_DIR)

        st.success(f"✔ تم استخراج {len(records)} جزء نصي!")

        texts = [rec["text"] for rec in records]
        ids = [rec["id"] for rec in records]

        st.info("جاري إنشاء Embeddings + FAISS...")
        index, model, embs = build_faiss_index(texts)

        st.session_state["records"] = records
        st.session_state["texts"] = texts
        st.session_state["ids"] = ids
        st.session_state["index"] = index
        st.session_state["model"] = model

# CHAT SECTION
st.subheader("اسأل...")

query = st.text_input("اكتب هنا…")

if st.button("إرسال"):
    if "index" not in st.session_state:
        st.error("يجب بناء قاعدة المعرفة أولاً!")
    else:
        model = st.session_state["model"]
        index = st.session_state["index"]
        ids = st.session_state["ids"]
        texts = st.session_state["texts"]

        results = search(query, model, index, ids, texts, k=8)
        answer = llm_answer(query, results)

        st.markdown("الإجابة")
        st.write(answer)

        st.markdown("---")
        st.markdown("الأجزاء المستخدمة")
        for r in results:
            st.write(f"**{r['id']}** — Score: {r['score']}")
            st.write(r["text"])
            st.write("---")


